"""Generate the Organizational Assessment worksheet and grade sheet.

Replaces the original 4.6 MB Word file, which carried embedded commercial fonts
and real people's names in its document metadata. These are built from scratch:
no embedded fonts, no author metadata.
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
OUT  = os.path.abspath(os.path.join(HERE, "..", "assignments", "files"))
TERM = "Fall 2026"

SECTIONS = [
    ("1. Organizational Culture", 2, "Maximum 250 words",
     ["Insert the photo of your follow-up meeting here:",
      "Briefly summarize how your client described their organization's culture:",
      "Would you want to work in a similar environment?"]),
    ("2. Website Analysis", 1, "Maximum 250 words",
     ["Comment on whether the site meets each of the seven bare essentials. "
      "Type out the bare essential, then start your comment after a dash."]),
    ("3. Bare Essentials for Your Organization", 2, "Maximum 250 words",
     ["Select the five most important bare essentials for your agency. For each, "
      "explain your client's actions or strategies in 1-2 sentences."]),
    ("4. Agency Goal", 2, "Maximum 250 words",
     ["Insert a screenshot of a document or web page describing the agency mission:",
      "Mission statement (quoted, with APA citation):",
      "An output of your agency:", "An outcome related to that output:",
      "Potential threats to accuracy in your project:",
      "Potential threats to representativeness in your project:"]),
    ("5. Describe Your Project", 1, "",
     ["Name of organization:", "Brief project title:",
      "Source of data (survey, records, other):", "Unit of analysis:",
      "Target population (include a number if possible):"]),
    ("6. Research Questions", 2, "",
     ["List 7-10 questions your agency hopes to answer with the data in your report:"]),
]

DEDUCTIONS = [("Lateness", "1 per day"), ("Spelling/Writing (WRT)", "Up to 3"),
              ("Organization/Neatness (ORG)", "Up to 5"),
              ("Project Scoping/Timeline", "")]

def base(title):
    d = Document()
    st = d.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(12)
    cp = d.core_properties
    cp.author = "PST 315"; cp.last_modified_by = "PST 315"
    cp.title = title; cp.comments = ""; cp.category = ""; cp.company = ""
    return d

def heading(d, text, size=14):
    p = d.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(size)
    return p

def fields(d, labels):
    for lab in labels:
        p = d.add_paragraph(); p.add_run(f"{lab}  ").bold = True
        p.add_run("_" * max(6, 62 - len(lab)))

# ---------------------------------------------------------------- worksheet
d = base("PST 315 Organizational Assessment")
h = heading(d, "PST 315 Organizational Assessment", 16); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = d.add_paragraph(TERM); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fields(d, ["Name:", "Date:", "Client and Agency Name:"])
d.add_paragraph()
for title, pts, limit, prompts in SECTIONS:
    heading(d, f"{title}  ({limit + ' - ' if limit else ''}{pts} point{'s' if pts != 1 else ''})")
    for pr in prompts:
        d.add_paragraph(pr)
        d.add_paragraph()
    d.add_paragraph()
d.save(os.path.join(OUT, "organizational-assessment.docx"))
print("  wrote organizational-assessment.docx")

# -------------------------------------------------------------- grade sheet
d = base("PST 315 Organizational Assessment Grade Sheet")
h = heading(d, "PST 315 Organizational Assessment Grade Sheet", 16)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = d.add_paragraph(TERM); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fields(d, ["Name:", "Date:", "Client and Agency Name:"])
d.add_paragraph()

def table(d, header, rows, total_label, total_val):
    t = d.add_table(rows=1, cols=3); t.style = "Table Grid"
    for i, x in enumerate(header):
        c = t.rows[0].cells[i]; c.text = ""; r = c.paragraphs[0].add_run(x); r.bold = True
    for name, mx in rows:
        cells = t.add_row().cells
        cells[0].text = name; cells[1].text = str(mx); cells[2].text = ""
    cells = t.add_row().cells
    cells[0].paragraphs[0].add_run(total_label).bold = True
    cells[1].paragraphs[0].add_run(str(total_val)).bold = True
    cells[2].text = ""
    return t

heading(d, "Points Earned", 12)
table(d, ["Section", "Maximum", "Actual"],
      [(s[0], s[1]) for s in SECTIONS], "Total Points Gained", 10)
d.add_paragraph()
heading(d, "Points Lost", 12)
table(d, ["Category", "Maximum", "Actual"], DEDUCTIONS, "Total Points Lost", "")
d.add_paragraph()
p = d.add_paragraph(); p.add_run("Points for Paper:  ").bold = True; p.add_run("_" * 20)
d.save(os.path.join(OUT, "organizational-assessment-grade-sheet.docx"))
print("  wrote organizational-assessment-grade-sheet.docx")
